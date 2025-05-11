from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT, WD_TABLE_ALIGNMENT
import matplotlib.pyplot as plt
import logging
import time
import os
from collections import defaultdict

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generates comprehensive Word document reports for API test results"""
    
    def __init__(self, test_errors, false_positives, response_times, test_result, 
                 test_statuses, start_time, end_time, base_url, env_info):
        """
        Initialize report generator with test data
        
        Args:
            test_errors (list): List of error messages from failed tests
            false_positives (list): track false positive error tests
            response_times (list): List of API response times
            test_result (unittest.TestResult): Test execution results
            test_statuses (list): list of the tests and it's statuses(passes and fails)
            start_time (float): Test execution start timestamp
            end_time (float): Test execution end timestamp
            base_url (str): Base API URL tested
            env_info (dict): Environment information
        """
        self.test_errors = test_errors
        self.false_positives = false_positives
        self.response_times = response_times
        self.test_result = test_result
        self.test_statuses = test_statuses
        self.start_time = start_time
        self.end_time = end_time
        self.base_url = base_url
        self.env_info = env_info
        self.doc = None  # Will hold the Word document object

        # Calculate test metrics
        self.total_tests = test_result.testsRun if test_result else 0
        self.passed = self.total_tests - len(test_result.failures) - len(test_result.errors) if test_result else 0
        self.failed = len(test_result.failures) + len(test_result.errors) if test_result else len(test_errors)

    def generate(self):
        """Generate the complete report document with all sections"""
        self._create_base_document()  # Initialize document structure
        self._add_summary_table()     # Add test summary table
        self._add_summary_chart()     # Add pie chart visualization
        self._analyze_failures() # Add tests passed and failed charts diving the errors
        self._add_response_time_chart()  # Add response time analysis
        self._add_response_time_stats() # avarage time, max and min time
        self._add_test_case_list()      # List Of All the Tests wish passes and fails
        self._add_environment_info()  # Add environment details
        self._add_execution_info()    # Add timing information
        self._add_error_section()     # Add detailed error reports

    def save(self, filename):
        """Save the generated document to file"""
        self.doc.save(filename)

    def _create_base_document(self):
        """Create the basic document structure with header and footer"""
        self.doc = Document()
        self._add_header()  # Add report title
        self._add_footer()  # Add page numbering


    def _add_header(self):
        """Add styled document header with title and timestamp"""
        # Main report title
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True
        
        # Title
        title = self.doc.add_paragraph()
        title_run = title.add_run("AUTOMATED TEST REPORT\n")
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Dark blue
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle_run = subtitle.add_run("Quality Assurance Department\n\n")
        subtitle_run.font.size = Pt(18)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Decorative line
        line = self.doc.add_paragraph()
        line_run = line.add_run()
        line_run.add_break()
        line.border_bottom = True

        # Add logo
        try:
            self.doc.add_picture('logo.png', width=Inches(2))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except FileNotFoundError:
            logger.warning("Company logo not found, skipping cover page image")



        # Project metadata
        meta = self.doc.add_paragraph()
        meta.add_run(f"\nProject: {self.env_info.get('project_name', 'N/A')}\n")
        meta.add_run(f"Environment: {self.env_info.get('environment', 'Staging')}\n")
        meta.add_run(f"Test Cycle: {self.env_info.get('test_cycle', 'Regression')}\n")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.runs[0].font.size = Pt(14)


        # Report Date
        date_str = time.strftime('%B %d, %Y %H:%M:%S', time.localtime() )
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(f"Generated on: {date_str}")
        date_run.italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break to content
        self.doc.add_page_break()

    def _add_footer(self):
        """Add page number footer to document"""
        section = self.doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right-align footer
        
        # Create page number field using Word XML
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    def _add_summary_table(self):
        """Add table showing test pass/fail summary"""
        table = self.doc.add_table(rows=2, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row with styling
        hdr_cells = table.rows[0].cells
        headers = [
            ('Total Tests', ''),
            ('Passed', '4CAF50'),  # Green
            ('Failed', 'FF5252'),   # Red
            ('Pass Rate', '2196F3'), # Blue
            ('False Positives', 'FFA500') # Orange
        ]
        
        for i, (text, color) in enumerate(headers):
            hdr_cells[i].text = text
            if color:
                hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Data row with conditional formatting
        row_cells = table.rows[1].cells
        
        # Total Tests
        row_cells[0].text = str(self.total_tests)
        
        # Passed (green)
        passed_cell = row_cells[1]
        passed_cell.text = str(self.passed)
        passed_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Pure Green (success color)
        
        # Failed (red with conditional intensity)
        failed_cell = row_cells[2]
        failed_cell.text = str(self.failed)
        fail_color = RGBColor(0xFF, 0x00, 0x00) if self.failed > 0 else RGBColor(0x99, 0x00, 0x00) #  Bright Red (when failures > 0)  Darker Red (when no failures)
        failed_cell.paragraphs[0].runs[0].font.color.rgb = fail_color
        
        # Pass Rate (percentage with icon)
        pass_rate = (self.passed / self.total_tests) * 100 if self.total_tests > 0 else 0
        pass_rate_cell = row_cells[3]
        pass_rate_cell.text = f"{pass_rate:.1f}%"
        
        # Color based on pass rate threshold
        if pass_rate >= 90:
            pass_rate_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Green (excellent performance)
        elif pass_rate >= 70:
            pass_rate_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xA5, 0x00) # Orange (warning/needs review)
            pass_rate_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red (critical condition)
        
        # False Positives
        fp_count = len(self.false_positives)
        fp_cell = row_cells[4]
        fp_cell.text = str(fp_count)
        if fp_count > 0:
            fp_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)  # Orange
        
        # Add contextual note
        if fp_count > 0:
            note = self.doc.add_paragraph()
            note.add_run("Note: ").bold = True
            note.add_run(f"{fp_count} test(s) returned 200 status but failed assertions")
            note.runs[0].font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
            note.style = 'Intense Quote'
        
        # Add visual separator
        self.doc.add_paragraph().add_run().add_break()

    def _add_summary_chart(self):
        """Generate and insert pie chart showing test results"""
        try:
            # Chart data
            labels = ['Passed', 'Failed']
            sizes = [self.passed, self.failed]
            colors = ['#4CAF50', '#FF5252']  # Green and red
            
            # Create pie chart
            fig, ax = plt.subplots()
            ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
            ax.axis('equal')  # Equal aspect ratio ensures circular pie
            
            # Save chart to temporary file
            chart_path = "test_summary_chart.png"
            plt.savefig(chart_path, bbox_inches='tight')
            plt.close(fig)
            
            # Insert chart into document
            self.doc.add_paragraph("Test Results Overview:")
            self.doc.paragraphs[-1].style = self.doc.styles['Heading 2']
            self.doc.add_picture(chart_path, width=Pt(300))  # 300 points wide
            os.remove(chart_path)  # Clean up temporary file
        except Exception as e:
            logger.error(f"Failed to generate chart: {str(e)}")

    def _analyze_failures(self):
        """Categorize failures and generate a visualization chart"""
        error_types = defaultdict(int)
        
        # Classify errors
        for error in self.test_errors:
            if "400" in str(error):
                error_types["Bad Request (400)"] += 1
            elif "401" in str(error):
                error_types["Unauthorized (401)"] += 1
            elif "404" in str(error):
                error_types["Not Found (404)"] += 1
            elif "500" in str(error):
                error_types["Server Error (500)"] += 1
            elif "Timeout" in str(error):
                error_types["Timeout"] += 1
            else:
                error_types["Other Errors"] += 1
        
        # Add false positives if they exist
        if hasattr(self, 'false_positives') and self.false_positives:
            error_types["False Positives (200)"] = len(self.false_positives)
        
        
        # Always show the table (even if empty)
        self._display_failure_table(error_types)

        # Generate the chart only if there are failures
        if error_types:
            self._generate_failure_chart(error_types)

    def _generate_failure_chart(self, error_types):
        """Generate a pie/donut chart of failure distribution"""
        try:
            import matplotlib.pyplot as plt
            
            # Prepare data
            labels = list(error_types.keys())
            sizes = list(error_types.values())
            colors = ['#FF5252', '#FFA500', '#FFD700', '#FF6347', '#9370DB', '#69B4FF']
            
            # Create figure
            fig, ax = plt.subplots(figsize=(8, 6))
            
            # Use a donut chart for better readability
            wedges, texts, autotexts = ax.pie(
                sizes, 
                labels=labels, 
                colors=colors,
                autopct='%1.1f%%',
                startangle=90,
                wedgeprops={'width': 0.4},  # Makes it a donut
                textprops={'fontsize': 8}
            )
            
            # Equal aspect ratio ensures the pie is circular
            ax.axis('equal')  
            plt.title('Failure Type Distribution', pad=20)
            
            # Save to temporary file
            chart_path = "failure_analysis_chart.png"
            plt.savefig(chart_path, bbox_inches='tight', dpi=150)
            plt.close(fig)
            
            # Insert into document
            self.doc.add_paragraph("Failure Analysis Chart:", style="Heading 3")
            self.doc.add_picture(chart_path, width=Pt(400))  # Adjust width as needed
            
            # Clean up
            os.remove(chart_path)
            
        except Exception as e:
            logger.error(f"Failed to generate failure chart: {str(e)}")
            self.doc.add_paragraph(
                f"⚠ Could not generate chart: {str(e)}", 
                style="Intense Quote"
            )

    def _display_failure_table(self, error_types):
        """Display failure statistics in a table"""
        self.doc.add_paragraph("Failure Statistics:", style="Heading 3")
        
        table = self.doc.add_table(rows=len(error_types)+1, cols=2)
        table.style = "Light Grid Accent 1"
        
        # Header
        table.cell(0, 0).text = "Error Type"
        table.cell(0, 1).text = "Count"
        
        # Data rows
        for i, (error_type, count) in enumerate(error_types.items(), 1):
            table.cell(i, 0).text = error_type
            table.cell(i, 1).text = str(count)
            
            # Highlight false positives
            if "False Positives" in error_type:
                table.cell(i, 1).font.color.rgb = RGBColor(255, 165, 0)  # Orange



    def _add_response_time_chart(self):
        """Generate and insert histogram of API response times grouped by endpoint"""
        try:
            if not self.response_times:
                return  # Skip if no response time data

            # Group response times by endpoint
            from collections import defaultdict
            endpoint_times = defaultdict(list)
            for entry in self.response_times:
                endpoint_times[entry['endpoint']].append(entry['duration'])

            # Create figure
            plt.figure(figsize=(12, 6))
            
            # Plot each endpoint's response times
            colors = ['#4CAF50', '#2196F3', '#FF5722', '#9C27B0']
            for i, (endpoint, times) in enumerate(endpoint_times.items()):
                plt.plot(times, marker='o', linestyle='-', 
                        color=colors[i % len(colors)], 
                        label=f"{endpoint} ({len(times)} calls)")

            plt.xlabel('Request Sequence')
            plt.ylabel('Response Time (seconds)')
            plt.title('Endpoint Response Times Over Test Execution')
            plt.legend()
            plt.grid(True)

            # Save and insert chart
            chart_path = "response_time_chart.png"
            plt.savefig(chart_path, bbox_inches='tight')
            plt.close()
            
            self.doc.add_paragraph("Endpoint Response Time Analysis:")
            self.doc.paragraphs[-1].style = self.doc.styles['Heading 2']
            self.doc.add_picture(chart_path, width=Pt(400))
            os.remove(chart_path)
        except Exception as e:
            logger.error(f"Failed to generate response time chart: {str(e)}")
            raise  # Re-raise to see error in test output


    def _add_test_case_list(self):
        """Add detailed table of all test cases with status"""
        self.doc.add_heading('Detailed Test Cases', level=1)
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header
        hdr = table.rows[0].cells
        hdr[0].text = 'Test Case ID'
        hdr[1].text = 'Test Name'
        hdr[2].text = 'Status'
        hdr[3].text = 'Duration (s)'
        # Populate data (assuming test_statuses is passed to ReportGenerator)
        for test in self.test_statuses:
            row = table.add_row().cells
            row[0].text = test['id']
            row[1].text = test['name']
            
            status = row[2].paragraphs[0].add_run(test['status'])
            status.bold = True
        
            # Color coding
            if test['status'] == 'Failed':
                status.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                status.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            
            # Always show duration, formatted to 3 decimal places
            row[3].text = f"{test['duration']:.3f}"

    def _add_response_time_stats(self):
        """Add table with response time statistics"""
        if not self.response_times:
            return

        stats = {
            'Average': sum(rt['duration'] for rt in self.response_times)/len(self.response_times),
            'Max': max(rt['duration'] for rt in self.response_times),
            'Min': min(rt['duration'] for rt in self.response_times)
        }

        self.doc.add_heading('Response Time Statistics', level=2)
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(stats.items()):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = f"{value:.2f} sec"
            

    def _add_environment_info(self):
        """Add table showing test environment details"""
        self.doc.add_heading('Environment Information', level=1)
        env_table = self.doc.add_table(rows=6, cols=2)
        env_table.style = self.doc.styles['Light List Accent 1']
        
        # Environment data to display
        env_info = [
            ("Base API URL", self.base_url),
            ("Python Version", self.env_info['python_version']),
            ("Platform", self.env_info['platform']),
            ("Requests Version", self.env_info['requests_version']),
            ("Hostname", self.env_info['hostname']),
            ("CPU Cores", self.env_info['cpu_cores']),
        ]
        
        # Populate table
        for i, (label, value) in enumerate(env_info):
            env_table.rows[i].cells[0].text = label
            env_table.rows[i].cells[1].text = str(value)
            env_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True  # Bold labels

    def _add_execution_info(self):
        """Add section with test execution timing information"""
        if self.start_time:
            total_duration = self.end_time - self.start_time
            minutes, seconds = divmod(total_duration, 60)
            
            self.doc.add_heading('Execution Information', level=1)
            exec_table = self.doc.add_table(rows=3, cols=2)
            exec_table.style = self.doc.styles['Light List Accent 1']
            
            # Start time row
            exec_table.rows[0].cells[0].text = "Test Execution Started"
            exec_table.rows[0].cells[1].text = time.strftime('%Y-%m-%d %H:%M:%S', 
                                                          time.localtime(self.start_time))
            
            # End time row
            exec_table.rows[1].cells[0].text = "Test Execution Finished"
            exec_table.rows[1].cells[1].text = time.strftime('%Y-%m-%d %H:%M:%S', 
                                                           time.localtime(self.end_time))
            
            # Duration row
            exec_table.rows[2].cells[0].text = "Total Test Duration"
            exec_table.rows[2].cells[1].text = f"{int(minutes)} minutes {seconds:.2f} seconds"
            
            # Bold all labels
            for row in exec_table.rows:
                row.cells[0].paragraphs[0].runs[0].bold = True

    def _add_error_section(self):
        """Add detailed error reports section"""
        self.doc.add_page_break()  # Start new page for errors
        self.doc.add_heading('Test Errors Report', level=1)
        self.doc.add_paragraph(f"The following {len(self.test_errors)} test(s) encountered errors:")
        self.doc.paragraphs[-1].style = self.doc.styles['Intense Quote']
        
        # Set monospace font for error output
        style = self.doc.styles['Normal']
        style.font.name = 'Consolas'
        style.font.size = Pt(10)
        
        # Add each error with formatting
        for error in self.test_errors:
            self._add_single_error_entry(error)
            self.doc.add_paragraph()  # Add spacing between errors

    def _add_single_error_entry(self, error):
        """Format and add a single error entry with syntax highlighting"""
        error_lines = error.strip().split("\n")
        for line in error_lines:
            paragraph = self.doc.add_paragraph()
            
            # Apply different formatting based on line content
            if line.startswith("Test Description:"):
                run = paragraph.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x8B)  # Dark blue
                
            elif line.startswith("Test:"):
                run = paragraph.add_run(line)
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue
                
            elif line.startswith("Error Type:"):
                run = paragraph.add_run(line)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red
                
            elif line.startswith("Error Message:"):
                run = paragraph.add_run(line)
                run.italic = True
                run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # Dark red
                self._add_error_message_border(paragraph)  # Add decorative border
                
            elif line.startswith("Request Body:"):
                run = paragraph.add_run(line)
                run.bold = True
                run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)  # Purple
                
            elif line.startswith("Response Status:") or line.startswith("Response URL:"):
                run = paragraph.add_run(line)
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green
                
            elif line.startswith("Response Content:"):
                run = paragraph.add_run(line)
                run.underline = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x80)  # Teal
                
            elif line.startswith("----------------------------------------"):
                self.doc.add_paragraph()  # Add extra spacing for separators
                
            else:
                # Color JSON content differently
                if line.strip().startswith(("{", "[")):
                    run = paragraph.add_run(line)
                    run.font.color.rgb = RGBColor(0xA5, 0x2A, 0x2A)  # Brown
                else:
                    paragraph.add_run(line)  # Default formatting

    def _add_error_message_border(self, paragraph):
        """Add decorative border around error messages"""
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement('w:pBorders')
        
        # Add borders to all sides
        for border_side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_side}')
            border.set(qn('w:val'), 'single')  # Single line
            border.set(qn('w:sz'), '12')       # Size
            border.set(qn('w:space'), '0')     # No spacing
            border.set(qn('w:color'), '8B0000')  # Dark red color
            borders.append(border)
        
        p_pr.append(borders)